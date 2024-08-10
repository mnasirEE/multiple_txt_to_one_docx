import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog
import os

# Step 1: Function to extract hostname from content
def extract_hostname(content):
    # Regex to find hostname, customize if needed
    match = re.search(r'hostname\s+(\S+)', content)
    if match:
        return match.group(1)
    return 'Unknown'

# Step 2: Function to extract site name from hostname
def extract_site_name(hostname):
    # Extract part before the first underscore
    parts = hostname.split('_')
    if parts:
        return parts[0]
    return 'Unknown'

# Step 3: Function to get the next appendix letter
def get_next_appendix_letter(current_letter):
    # Function to convert a letter (e.g., 'A') to its corresponding position (e.g., 0 for 'A')
    def letter_to_number(letter):
        return ord(letter) - ord('A')

    # Function to convert a number (e.g., 0) to its corresponding letter (e.g., 'A')
    def number_to_letter(number):
        return chr(number + ord('A'))

    # Convert current letter to a list of its positions in the alphabet (e.g., 'AB' -> [0, 1])
    positions = [letter_to_number(c) for c in current_letter]

    # Increment the letter positions
    for i in reversed(range(len(positions))):
        if positions[i] < 25:  # 25 corresponds to 'Z'
            positions[i] += 1
            break
        positions[i] = 0  # Reset current position to 'A' if it overflows to 'Z'

    # If all positions have overflowed (e.g., from 'Z' -> 'AA'), add a new letter
    if all(p == 0 for p in positions):
        positions.insert(0, 0)

    # Convert positions back to letters and return as a string
    next_letter = ''.join(number_to_letter(p) for p in positions)
    return next_letter


# Step 4: Function to add title and content from each file to the document
def add_file_to_document(doc, file_path, appendix_letter, separator_style):
    with open(file_path, 'r') as file:
        content = file.read()
        
        # Step 6: Search for the specific string
        if 'no ip http server' not in content:
            print(f"Error: 'no ip http server' not found in {file_path}")

        # Step 1: Extract hostname from content
        hostname = extract_hostname(content)
        
        # Extract site name from the hostname
        site_name = extract_site_name(hostname)
        
        # Add title string with appendix letter
        title_paragraph = doc.add_paragraph()
        run = title_paragraph.add_run(f"Appendix {appendix_letter} Configuration for {site_name}: {hostname}")
        run.bold = True
        run.font.size = Pt(separator_style['size'])
        run.font.name = separator_style['font']
        title_paragraph.style.font.name = separator_style['font']
        title_paragraph.style.element.rPr.rFonts.set(qn('w:eastAsia'), separator_style['font'])
        
        # Add content to the document
        doc.add_paragraph(content, style='Normal')
        
        # Add spacing between sections
        doc.add_paragraph()

# Step 5: Function to combine all txt files into a single docx
def combine_txt_files_to_docx(txt_files, output_docx, separator_style):
    doc = Document()
    doc.styles['Normal'].font.size = Pt(8)  # Set body text font size
    
    appendix_letter = 'D'  # Start from Appendix D
    for file_path in txt_files:
        add_file_to_document(doc, file_path, appendix_letter, separator_style)
        # Update appendix letter for the next file
        appendix_letter = get_next_appendix_letter(appendix_letter)
    
    doc.save(output_docx)

# Step 6: Function to open file dialog and select files
def select_files():
    root = tk.Tk()
    root.withdraw()  # Hide the root window
    file_paths = filedialog.askopenfilenames(
        title='Select text files',
        filetypes=[('Text files', '*.txt')]
    )
    return list(file_paths)

# Step 7: Main function to run the script
if __name__ == "__main__":
    txt_files = select_files()  # Select files through dialog
    if not txt_files:
        print("No files selected.")
    else:
        output_docx = 'combined_document.docx'
        separator_style = {'font': 'Calibri', 'size': 16}
        combine_txt_files_to_docx(txt_files, output_docx, separator_style)
        print(f"Document saved as {output_docx}")
